from datetime import datetime

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from fastapi import APIRouter
from fastapi.responses import FileResponse
from sqlalchemy import insert, select

from src.api.dependencies import UserIdDep
from src.database import async_session_maker
from src.models.basType import BASTypes
from src.models.orders import Orders

router = APIRouter(prefix="/order", tags=["Заказы"])
#
# import httpx
#
# async def get_prediction_from_ml():
#     async with httpx.AsyncClient() as client:
#         response = await client.post("http://localhost:8001/predict")
#         return response.json()
#
# @router.get("/predict")
# async def get_prediction():
#     return await get_prediction_from_ml()
#
# @router.get("/prompt/")
# def get_prompt(prompt: str):
#     ...

@router.post("/prompt/")
async def add_prompt(text: str, user_id: UserIdDep):
    async with async_session_maker() as session:
        add_stmt = insert(Orders).values(user_id=user_id, prompt=text).returning(Orders)
        res = await session.execute(add_stmt)
        prompt = res.scalars().one()
        await session.commit()
    return prompt

@router.get("/prompt/")
async def get_prompts(user_id: UserIdDep):
    async with async_session_maker() as session:
        add_stmt = select(Orders).filter_by(user_id=user_id)
        res = await session.execute(add_stmt)
        prompts = res.scalars().all()
        await session.commit()
    return {"data": prompts}

@router.post("/")
async def load():
    async with async_session_maker() as session:
        add_stmt = insert(BASTypes).values(
            name="Сельское хозяйство",
            type="""Мультиспектральные\n
Опрыскивающие дроны
""",
            tasks="""Мониторинг состояния посевов\n
Анализ вегетационных индексов (NDVI, NDRE)\n
Точечное внесение удобрений и пестицидов\n
Контроль орошения\n
Инвентаризация земель\n
Оценка ущерба (град, засуха, подтопление).
        """,
            equipments="""Мультиспектральные камеры (NDVI, GNDVI, RedEdge):\n
Анализ здоровья растений.\n
Выявление стрессовых участков.\n
Тепловизоры — для оценки водного стресса.\n
RGB-камеры высокого разрешения — для ортофото, документации.\n
(опционально) Гиперспектральные камеры — для точных научных задач.
""",
            fly="""Время полёта: 40–55 мин (в большинстве случаев хватает).\n
Автоматические миссии (облёт объекта по кругу, сетка).\n
Дальность связи: до 5–15 км.\n
Быстрая подготовка к полёту — критично на стройплощадке.\n
Грузоподъемность: 2-6 кг (для лидаров и профессиональных камер)""",
            navigation="""
                    RTK/PPK — важно для точного позиционирования и повторяемости миссий.\n
Автоматические маршруты (по границам поля)\n
Датчики высоты (для работы на неровном рельефе)
        """,
            conditions="""
                    Работа при ветре до 10 м/с.\n
Температура: от 0°C и выше.\n
Пыле- и влагозащита: IP54–67 (особенно для опрыскивающих дронов).\n
устойчивость к химикатам\n
Работа вдали от связи — важна автономность и оффлайн-обработка.
""",
            software="""Распылительное оборудование:\n
Бак: 20-40 л\n
Ширина захвата: 6-8 м\n
Точность внесения: ±5%\n
Дополнительные датчики:\n
Датчик влажности почвы\n
Анемометр (замер ветра\n
        """
        )
        await session.execute(add_stmt)
        await session.commit()
        return {"status": "Ok"}


async def get_type(name: str):
    async with async_session_maker() as session:
        query = select(BASTypes).filter_by(name=name)
        res = await session.execute(query)
        type = res.scalars().one()
        await upload_to_file(type)
        return type

async def upload_to_file(type):
    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    title = doc.add_heading('Отчет по дрону', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"Дата генерации: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("\n")

    p = doc.add_paragraph(type.name)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.type)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.tasks)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.equipments)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.fly)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.navigation)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.conditions)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph(type.software)
    p.paragraph_format.space_after = Pt(6)

    file_path="/app/file.docx"
    doc.save(file_path)
    return FileResponse(
        path=file_path,
        filename="file.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
